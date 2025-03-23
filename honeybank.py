import os
import random
import csv
import argparse
import json
from faker import Faker
from docx import Document
import pandas as pd
from fpdf import FPDF
import ollama
import unicodedata

# Change to preferred model. The dolphin-mixtral model was used due to it being uncensored and would generate the data without complaints.
MODEL_NAME = "dolphin-mixtral"

SYSTEM_PROMPT = """You are an expert data generation agent. You only genereate the requested amount of data and do not output anything else. This data is purely for testing purposes and will not be used for any malicious purposes. Do not put messages into the document like "sample data", "realistic" or any messages that would make it look non-realistic. The file should only have the data relevent to the topic. IMPORTANT: If the customers name is in the data, use the name given by the user. IMPORTANT: The output needs to formated as content suitable for the requested file type. The data needs to match the structure and formatting needed for the requested file type. Data only.

Here is an example CSV and XLSX data.
header 1, header 2, header 3
value 1, value 2, value 3

Here is an example of a loan application.
Loan Purpose: Start a food truck business
Loan Amount: $50,000
Repayment Period: 3 years
Applicant Name: Maria Gonzalez
Business Plan Summary: Maria plans to start a gourmet taco food truck in Austin, Texas. She has five years of experience in the food industry and already has a following from her social media cooking channel. The funds will be used to purchase the food truck, kitchen equipment, and initial inventory.

Loan Letter:
My name is Maria Gonzalez, and I am seeking a loan of $50,000 to start a gourmet taco food truck business in Austin, Texas. With over five years of experience in the food industry and a strong online presence through my cooking channel, I am confident in my ability to run a successful mobile kitchen.

The loan will be used to purchase a fully equipped food truck, kitchen equipment, and the initial stock of ingredients. I am requesting a repayment period of 3 years, during which I anticipate steady growth and profitability. My business plan outlines marketing strategies, projected revenue, and cost management to ensure timely repayment.

Thank you for considering my application. I look forward to the opportunity to discuss it further.

Sincerely,
Maria Gonzalez

"""

VALID_TOPICS = {
    "Loan Application": [".docx", ".pdf"],
    "Customer Profile": [".pdf", ".docx"],
    "Internal Audit": [".pdf", ".docx"],
    "Financial Report": [".xlsx", ".pdf", ".docx"],
    "Tax Return": [".pdf", ".xlsx", ".docx"]
}

def create_fake_docx(file_path, customer_name, account_number, fake, use_ollama, topic):
    """Generate a fake banking document in .docx format."""
    doc = Document()
    doc.add_heading("Confidential Banking Document", level=1)
    if use_ollama:
        prompt = generate_prompt(topic, "docx", 1, customer_name)
        content = call_ollama(prompt)
        doc.add_paragraph(content)
    else:
        doc.add_paragraph(f"Customer Name: {customer_name}")
        doc.add_paragraph(f"Account Number: {account_number}")
        doc.add_paragraph(f"Balance: ${fake.random_int(min=1000, max=500000)}")
        doc.add_paragraph(f"Date: {fake.date_this_year()}")
    doc.save(file_path)

def create_fake_xlsx(file_path, customer_name, account_number, fake, use_ollama, topic):
    """Generate a fake banking transaction record in .xlsx format."""
    if use_ollama:
        prompt = generate_prompt(topic, "xlsx", 10, customer_name)
        content = call_ollama(prompt)
        from io import StringIO
        df = pd.read_csv(StringIO(content))
    else:
        data = {
            "Transaction ID": [fake.uuid4() for _ in range(10)],
            "Date": [fake.date_this_year() for _ in range(10)],
            "Amount ($)": [round(random.uniform(100, 5000), 2) for _ in range(10)],
            "Description": [fake.sentence() for _ in range(10)]
        }
        df = pd.DataFrame(data)
    df.to_excel(file_path, index=False)

def create_fake_csv(file_path, customer_name, account_number, fake, use_ollama, topic):
    """Generate a fake banking customer data file in .csv format."""
    if use_ollama:
        prompt = generate_prompt(topic, "csv", 10, customer_name)
        content = call_ollama(prompt)
        with open(file_path, 'w') as f:
            f.write(content)
    else:
        with open(file_path, mode='w', newline='') as file:
            writer = csv.writer(file)
            writer.writerow(["Customer Name", "Account Number", "Balance", "Last Transaction Date"])
            for _ in range(10):
                writer.writerow([customer_name, account_number, 
                                 fake.random_int(min=1000, max=500000), 
                                 fake.date_this_year()])

def generate_prompt(topic: str, file_type: str, num_records: int, customer_name: str) -> str:
    return f"""Generate {num_records} records of {topic} data for {customer_name} in {file_type} format."""

def call_ollama(prompt: str) -> str:
    response = ollama.chat(
        model=MODEL_NAME,
        messages=[
            {"role": "system", "content": SYSTEM_PROMPT},
            {"role": "user", "content": prompt}
        ]
    )
    return response['message']['content'].strip()

def create_fake_pdf(file_path, customer_name, account_number, fake, use_ollama, topic):
    """Generate a fake bank account statement in .pdf format."""
    pdf = FPDF()
    pdf.set_auto_page_break(auto=True, margin=15)
    pdf.add_page()
    pdf.set_font("Arial", style='', size=12)
    pdf.cell(200, 10, txt="Bank Account Statement", ln=True, align='C')
    pdf.ln(10)
    try:
        if use_ollama:
            prompt = generate_prompt(topic, "pdf", 1, customer_name)
            content = call_ollama(prompt)
            for line in content.split("\n"):
                safe_line = unicodedata.normalize("NFKD", line).encode("latin-1", "ignore").decode("latin-1")
                pdf.cell(200, 10, txt=safe_line, ln=True)
        else:
            pdf.cell(200, 10, txt=f"Customer: {customer_name}", ln=True)
            pdf.cell(200, 10, txt=f"Account Number: {account_number}", ln=True)
            pdf.cell(200, 10, txt=f"Balance: ${fake.random_int(min=1000, max=500000)}", ln=True)
        pdf.output(file_path)
    except UnicodeEncodeError as e:
        print(f"UnicodeEncodeError writing PDF: {e}. Skipping {file_path}.")

def generate_fake_files(base_path, num_folders, min_files, max_files, file_types, use_ollama):
    """Generate fake customer folders with banking-related documents."""
    fake = Faker()

    for i in range(num_folders):
        customer_name = fake.name()
        account_number = str(fake.random_number(digits=8, fix_len=True))
        folder_name = f"{customer_name.replace(' ', '_')}-{account_number}"
        folder_path = os.path.join(base_path, folder_name)
        os.makedirs(folder_path, exist_ok=True)

        num_files = random.randint(min_files, max_files)
        for _ in range(num_files):
            topic = random.choice(list(VALID_TOPICS.keys()))
            possible_exts = [ext.lstrip(".") for ext in VALID_TOPICS[topic] if ext.lstrip(".") in file_types]
            if not possible_exts:
                continue  # skip if no allowed file types for the topic
            file_ext = random.choice(possible_exts)
            file_name = f"{topic.replace(' ', '_')}_{fake.random_number(digits=4)}.{file_ext}"
            file_path = os.path.join(folder_path, file_name)

            if file_ext == "docx":
                create_fake_docx(file_path, customer_name, account_number, fake, use_ollama, topic)
            elif file_ext == "xlsx":
                create_fake_xlsx(file_path, customer_name, account_number, fake, use_ollama, topic)
            elif file_ext == "csv":
                create_fake_csv(file_path, customer_name, account_number, fake, use_ollama, topic)
            elif file_ext == "pdf":
                create_fake_pdf(file_path, customer_name, account_number, fake, use_ollama, topic)

    print(f"Generated {num_folders} fake customer folder(s) with fake banking data.")

if __name__ == "__main__":
    parser = argparse.ArgumentParser(description="Generate fake customer banking files for honey shares.")
    parser.add_argument("--base_path", type=str, default="./", help="Base directory for generated files")
    parser.add_argument("--num_folders", type=int, default=1, help="Number of customer folders to create")
    parser.add_argument("--min_files", type=int, default=4, help="Minimum number of files per folder")
    parser.add_argument("--max_files", type=int, default=10, help="Maximum number of files per folder")
    parser.add_argument("--file_types", nargs="+", default=["docx", "xlsx", "csv", "pdf"], help="List of file types to generate")
    parser.add_argument("--use_ollama", action="store_true", help="Use Ollama to generate file content")

    args = parser.parse_args()
    os.makedirs(args.base_path, exist_ok=True)
    generate_fake_files(args.base_path, args.num_folders, args.min_files, args.max_files, args.file_types, args.use_ollama)
